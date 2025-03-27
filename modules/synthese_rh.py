# modules/synthese_rh.py

from docx import Document
import io

def generer_rapport_rh(df, agent_id, params):
    doc = Document()
    doc.add_heading(f"📋 Rapport RH – Agent {agent_id}", 0)

    agent_data = df[df["Agent"] == agent_id]
    if agent_data.empty:
        doc.add_paragraph("Aucune donnée disponible.")
        return

    doc.add_heading("📌 Écarts par KPI", level=1)
    for kpi in params["kpi"]:
        ecart = round(agent_data[f"Ecart_{kpi}"].mean() * 100, 2)
        doc.add_paragraph(f"{kpi} : {ecart}%", style="List Bullet")

    score = round(agent_data["Score_Global"].mean() * 100, 2)
    doc.add_heading("🎯 Score Global Moyen", level=1)
    doc.add_paragraph(f"{score}%")

    doc.add_heading("✍️ Commentaire RH", level=1)
    doc.add_paragraph("À compléter...")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
